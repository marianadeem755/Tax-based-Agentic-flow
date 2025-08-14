
# =========================================
# 📋 Pakistan Tax Form Finder - Enhanced Version
# =========================================
import streamlit as st
import os, json, time, base64, tempfile
from io import BytesIO
import requests
import pycountry
from bs4 import BeautifulSoup
from urllib.parse import urljoin, quote_plus
from dotenv import load_dotenv
from groq import Groq
import pandas as pd
import openpyxl
from openpyxl import load_workbook

# Fixed PyMuPDF import
try:
    import fitz  # PyMuPDF
except ImportError:
    st.error("PyMuPDF not available. PDF features will be limited.")
    fitz = None

# Document processing imports with fallbacks
try:
    import mammoth
except ImportError:
    mammoth = None
    
try:
    from docx import Document
except ImportError:
    Document = None

import re
# ─────────────────────────────────────────
# 🔐 Load API keys
# ─────────────────────────────────────────
load_dotenv()
SERPER_API_KEY = os.getenv("SERPER_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Initialize Groq client if API key is available
groq_client = None
if GROQ_API_KEY:
    groq_client = Groq(api_key=GROQ_API_KEY)

# ─────────────────────────────────────────
# 🧠 Enhanced Agent Classification
# ─────────────────────────────────────────
def classify_query_mode(query: str) -> int:
    """Enhanced classification with better prompting"""
    if not groq_client:
        # Better fallback logic
        tax_keywords = ['tax', 'form', 'filing', 'deduction', 'income', 'return', 'fbr', 'ntn', 'withholding']
        return 0 if any(keyword in query.lower() for keyword in tax_keywords) else 1
    
    prompt = f"""
    Classify this user query as either TAX ASSISTANT (0) or GENERAL CHATBOT (1).
    
    TAX ASSISTANT (0): Pakistani tax forms, filing procedures, tax calculations, deductions, FBR regulations, NTN requirements, income tax, sales tax, withholding tax, tax deadlines, tax compliance, form filling help.
    
    GENERAL CHATBOT (1): General conversation, non-tax questions, casual chat, personal topics, other countries' taxes, non-Pakistani tax systems.
    
    Query: "{query}"
    
    Answer with only 0 or 1:"""
    
    try:
        completion = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=1,
        )
        result = completion.choices[0].message.content.strip()
        return 0 if result.startswith("0") else 1
    except Exception:
        tax_keywords = ['tax', 'form', 'filing', 'deduction', 'income', 'return', 'fbr', 'ntn']
        return 0 if any(keyword in query.lower() for keyword in tax_keywords) else 1

def identify_tax_type(query: str) -> str:
    """Identify specific tax type from user query"""
    if not groq_client:
        # Fallback keyword matching
        query_lower = query.lower()
        if any(word in query_lower for word in ['income', 'salary', 'itr']):
            return 'income_tax'
        elif any(word in query_lower for word in ['sales', 'gst', 'sst']):
            return 'sales_tax'
        elif any(word in query_lower for word in ['withholding', 'advance', 'deduction']):
            return 'withholding_tax'
        elif any(word in query_lower for word in ['property', 'real estate', 'immovable']):
            return 'property_tax'
        elif any(word in query_lower for word in ['wealth', 'asset']):
            return 'wealth_tax'
        elif any(word in query_lower for word in ['business', 'company', 'corporate']):
            return 'business_tax'
        else:
            return 'general_tax'
    
    try:
        prompt = f"""
        Identify the specific tax type from this Pakistani tax query. Choose ONLY ONE:
        
        - income_tax: Income tax, salary tax, ITR, personal tax
        - sales_tax: Sales tax, GST, SST, value added tax
        - withholding_tax: Withholding tax, advance tax, tax deduction
        - property_tax: Property tax, real estate tax, immovable property
        - wealth_tax: Wealth tax, wealth statement, asset declaration
        - business_tax: Business tax, corporate tax, company registration
        - general_tax: General tax questions, multiple types, unclear
        
        Query: "{query}"
        
        Answer with only the tax type (e.g., income_tax):"""
        
        completion = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=20,
        )
        
        result = completion.choices[0].message.content.strip()
        valid_types = ['income_tax', 'sales_tax', 'withholding_tax', 'property_tax', 'wealth_tax', 'business_tax', 'general_tax']
        
        for tax_type in valid_types:
            if tax_type in result:
                return tax_type
        
        return 'general_tax'
        
    except Exception:
        return 'general_tax'

# ─────────────────────────────────────────
# 🌐 Enhanced Search helpers
# ─────────────────────────────────────────
def serper_search(query, country_code="pk"):
    """Enhanced search for Pakistani tax forms using Serper API - returns more results"""
    url = "https://google.serper.dev/search"
    
    # Enhanced search queries for better results
    search_queries = [
        f"{query} fillable tax form site:.gov.pk OR site:.fbr.gov.pk filetype:pdf",
        f"{query} tax form Pakistan FBR filetype:pdf",
        f"{query} form download Pakistan taxation filetype:pdf",
        f"Pakistan {query} form FBR official filetype:pdf"
    ]
    
    all_results = []
    
    try:
        if not SERPER_API_KEY:
            return fallback_search(query, country_code)
        
        headers = {"X-API-KEY": SERPER_API_KEY}
        
        for search_query in search_queries[:2]:  # Use first 2 queries to avoid rate limits
            data = {"q": search_query, "gl": "pk", "hl": "en", "num": 5}
            
            try:
                r = requests.post(url, json=data, headers=headers, timeout=15)
                if r.status_code == 200:
                    results = r.json().get("organic", [])
                    all_results.extend(results)
                time.sleep(0.5)  # Rate limiting
            except Exception:
                continue
        
        # Remove duplicates and return up to 8 results
        seen_urls = set()
        unique_results = []
        for result in all_results:
            url = result.get('link', '')
            if url not in seen_urls and len(unique_results) < 8:
                seen_urls.add(url)
                unique_results.append(result)
        
        return unique_results if unique_results else fallback_search(query, country_code)
        
    except Exception:
        return fallback_search(query, country_code)

def fallback_search(query, country_code=""):
    """Enhanced fallback search method with multiple sources"""
    all_results = []
    
    try:
        # DuckDuckGo search
        search_query = quote_plus(f"{query} Pakistan tax form PDF FBR")
        ddg_url = f"https://ddg-api.herokuapp.com/search?query={search_query}&limit=8"
        
        response = requests.get(ddg_url, timeout=10)
        if response.status_code == 200:
            results = response.json()
            for result in results:
                all_results.append({
                    "title": result.get("title", ""),
                    "link": result.get("link", ""),
                    "snippet": result.get("snippet", "")
                })
    
    except Exception:
        pass
    
    # Add some default Pakistani tax form sources if no results
    if len(all_results) < 3:
        default_sources = [
            {
                "title": "FBR Official Forms - Income Tax Returns",
                "link": "https://www.fbr.gov.pk/categ/income-tax-forms/50149/132/1",
                "snippet": "Official FBR income tax forms and returns for individuals and companies"
            },
            {
                "title": "Sales Tax Forms - FBR Pakistan", 
                "link": "https://www.fbr.gov.pk/categ/sales-tax-forms/50150/132/1",
                "snippet": "FBR sales tax registration and return forms"
            },
            {
                "title": "Withholding Tax Forms - Federal Board of Revenue",
                "link": "https://www.fbr.gov.pk/categ/withholding-tax-forms/50151/132/1", 
                "snippet": "Official withholding tax statements and forms"
            }
        ]
        all_results.extend(default_sources)
    
    return all_results

# ─────────────────────────────────────────
# 📄 Enhanced Document Processing
# ─────────────────────────────────────────
def is_pdf_fillable(file_bytesio):
    """Enhanced PDF fillability check"""
    try:
        file_bytesio.seek(0)
        doc = fitz.open(stream=file_bytesio, filetype="pdf")
        
        fillable_fields_count = 0
        for page in doc:
            widgets = page.widgets()
            if widgets:
                try:
                    widget_count = len(widgets)
                except TypeError:
                    widget_list = list(widgets)
                    widget_count = len(widget_list)
                fillable_fields_count += widget_count
        
        doc.close()
        return fillable_fields_count > 0, fillable_fields_count
    except Exception as e:
        st.error(f"Error checking PDF: {str(e)}")
        return False, 0

def analyze_excel_file(file_bytesio):
    """Enhanced Excel analysis"""
    try:
        file_bytesio.seek(0)
        df = pd.read_excel(file_bytesio, sheet_name=None)
        sheets_info = []
        total_fields = 0
        
        for sheet_name, sheet_df in df.items():
            empty_cells = sheet_df.isnull().sum().sum()
            form_keywords = ['name', 'amount', 'tax', 'income', 'deduction', 'total', 'date']
            form_fields = sum(1 for col in sheet_df.columns 
                            if any(keyword in str(col).lower() for keyword in form_keywords))
            
            sheets_info.append({
                'name': sheet_name,
                'rows': len(sheet_df),
                'columns': len(sheet_df.columns),
                'potential_fields': form_fields
            })
            total_fields += form_fields
        
        return {
            'type': 'excel',
            'sheets': sheets_info,
            'total_potential_fields': total_fields,
            'is_form_like': total_fields > 2
        }
    except Exception as e:
        st.error(f"Error analyzing Excel: {str(e)}")
        return None

def analyze_word_file(file_bytesio):
    """Enhanced Word analysis"""
    try:
        file_bytesio.seek(0)
        doc = Document(file_bytesio)
        
        form_fields = 0
        form_keywords = ['_____', '[ ]', 'name:', 'amount:', 'tax:', 'income:', 'date:']
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.lower()
            form_fields += sum(1 for keyword in form_keywords if keyword in text)
        
        return {
            'type': 'word',
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'form_fields_found': form_fields,
            'is_form_like': form_fields > 2 or len(doc.tables) > 1
        }
    except Exception as e:
        st.error(f"Error analyzing Word: {str(e)}")
        return None

def extract_form_fields(file_bytesio):
    """Enhanced form field extraction"""
    try:
        file_bytesio.seek(0)
        doc = fitz.open(stream=file_bytesio, filetype="pdf")
        
        fields = []
        widget_types = {
            fitz.PDF_WIDGET_TYPE_TEXT: "Text Field",
            fitz.PDF_WIDGET_TYPE_CHECKBOX: "Checkbox",
            fitz.PDF_WIDGET_TYPE_RADIOBUTTON: "Radio Button",
            fitz.PDF_WIDGET_TYPE_COMBOBOX: "Dropdown",
            fitz.PDF_WIDGET_TYPE_LISTBOX: "List Box"
        }
        
        for page_num, page in enumerate(doc):
            widgets = page.widgets()
            
            if widgets:
                try:
                    widget_list = list(widgets)
                except:
                    widget_list = []
                    for widget in widgets:
                        widget_list.append(widget)
                
                for widget in widget_list:
                    field_type = widget_types.get(widget.field_type, "Unknown")
                    field_info = {
                        "name": widget.field_name or f"Field_{page_num}_{len(fields)}",
                        "type": field_type,
                        "value": getattr(widget, 'field_value', ''),
                        "page": page_num + 1
                    }
                    fields.append(field_info)
        
        doc.close()
        return fields
    except Exception as e:
        st.error(f"Error extracting fields: {str(e)}")
        return []

def fetch_pdf(url):
    """Enhanced PDF fetching - downloads all PDFs regardless of fillable status"""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        
        r = requests.get(url, headers=headers, timeout=20, allow_redirects=True)
        
        if r.status_code == 200:
            # Check if it's a PDF
            content_type = r.headers.get('Content-Type', '').lower()
            if 'application/pdf' in content_type or url.lower().endswith('.pdf'):
                pdf_bytes = BytesIO(r.content)
                return pdf_bytes
            else:
                # Try to find PDF links in HTML
                soup = BeautifulSoup(r.text, "html.parser")
                pdf_links = []
                
                # Look for PDF links
                for link in soup.find_all('a', href=True):
                    href = link['href'].lower()
                    if href.endswith('.pdf') or 'pdf' in href:
                        pdf_url = link['href'] if link['href'].startswith("http") else urljoin(url, link['href'])
                        pdf_links.append(pdf_url)
                
                # Try to fetch the first PDF link found
                for pdf_link in pdf_links[:1]:  # Try only the first PDF link
                    try:
                        pdf_response = requests.get(pdf_link, headers=headers, timeout=15)
                        if pdf_response.status_code == 200:
                            return BytesIO(pdf_response.content)
                    except:
                        continue
        
        return None
    except Exception:
        return None

def display_pdf(file_bytesio):
    """Enhanced PDF display with download option"""
    try:
        file_bytesio.seek(0)
        base64_pdf = base64.b64encode(file_bytesio.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="600" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)
        
        file_bytesio.seek(0)
        st.download_button(
            label="📥 Download PDF",
            data=file_bytesio,
            file_name="pakistan_tax_form.pdf",
            mime="application/pdf"
        )
    except Exception as e:
        st.error(f"Error displaying PDF: {str(e)}")

# ─────────────────────────────────────────
# 📊 Enhanced Form Analysis using Groq LLM
# ─────────────────────────────────────────
def analyze_form_with_llm(file_type, form_fields=None, file_analysis=None, search_query=None):
    """Dynamic form analysis and recommendations using Groq LLM"""
    if not groq_client:
        return "LLM analysis requires Groq API key configuration."
    
    try:
        # Build context for analysis
        context = f"File Type: {file_type}\n"
        
        if search_query:
            context += f"Search Query: {search_query}\n"
            
        if file_type == 'pdf' and form_fields:
            field_names = [field['name'] for field in form_fields[:20]]
            field_types = [field['type'] for field in form_fields[:20]]
            context += f"PDF Fields ({len(form_fields)} total): {', '.join(field_names[:10])}\n"
            context += f"Field Types: {', '.join(set(field_types))}\n"
            
        elif file_analysis:
            if file_type == 'excel':
                context += f"Excel Sheets: {len(file_analysis.get('sheets', []))}\n"
                context += f"Potential Fields: {file_analysis.get('total_potential_fields', 0)}\n"
            elif file_type == 'word':
                context += f"Paragraphs: {file_analysis.get('paragraphs', 0)}\n"
                context += f"Tables: {file_analysis.get('tables', 0)}\n"
                context += f"Form Elements: {file_analysis.get('form_fields_found', 0)}\n"

        prompt = f"""
        You are a Pakistani tax expert analyzing a tax form. Based on the form details below, provide comprehensive analysis and recommendations.

        {context}

        Provide analysis in the following structure:
        
        🎯 **Form Identification**: Identify the likely tax form type and purpose
        📋 **Tax Category**: Determine which Pakistani tax category this belongs to (Income Tax, Sales Tax, Withholding Tax, etc.)
        📊 **Complexity Level**: Assess if this is basic, intermediate, or advanced level form
        📅 **Filing Timeline**: Typical deadlines and submission requirements for this type
        📝 **Key Sections**: Important sections or fields that require special attention
        ⚠️ **Common Issues**: Typical mistakes or challenges with this form type
        📚 **Required Documents**: Essential documents needed to complete this form
        💡 **Expert Tips**: Professional recommendations for accurate completion
        🔗 **FBR Requirements**: Specific FBR regulations or compliance notes
        📞 **Next Steps**: Recommended actions after analysis

        Keep recommendations specific to Pakistani tax law and FBR requirements.
        """

        completion = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1200,
        )

        return completion.choices[0].message.content

    except Exception as e:
        return f"Error in LLM analysis: {str(e)}"

# ─────────────────────────────────────────
# 🤖 Enhanced LLM Agents with Tax Type Specific Responses
# ─────────────────────────────────────────
def tax_agent_response(user_query, tax_form_type=None, form_fields=None):
    """Enhanced tax agent with comprehensive Pakistani tax knowledge"""
    if not groq_client:
        return "I need an LLM API key to provide detailed assistance."
    
    try:
        # Identify tax type from query
        tax_type = identify_tax_type(user_query)
        
        # Build context for the LLM
        context = f"""
        You are LifePilot Tax Agent - Pakistan's leading AI tax consultant with expertise in all FBR regulations and Pakistani taxation laws.
        
        USER QUERY: "{user_query}"
        IDENTIFIED TAX TYPE: {tax_type.replace('_', ' ').title()}
        """
        
        if tax_form_type:
            context += f"\nCURRENT FORM TYPE: {tax_form_type}"
            
        if form_fields and len(form_fields) > 0:
            fields_info = ", ".join([f["name"] for f in form_fields[:8]])
            context += f"\nFORM FIELDS AVAILABLE: {fields_info}"
        
        # Enhanced prompt for comprehensive response
        prompt = f"""
        {context}
        
        As LifePilot Tax Agent, provide a COMPREHENSIVE, EXPERT response for Pakistani tax matters.
        
        Structure your response with these sections as needed (use emojis and headers):
        
        🎯 **Direct Answer**: Clear, specific response to the user's question
        📋 **Relevant Forms**: Exact FBR form names/numbers needed (if applicable)  
        📊 **Step-by-Step Process**: Detailed instructions with sequence numbers
        📅 **Important Deadlines**: Current tax year deadlines (mention to verify dates)
        ⚠️ **Compliance Requirements**: Legal obligations and penalties to avoid
        💰 **Financial Calculations**: Tax rates, calculations, or savings tips (if applicable)
        📚 **Required Documents**: Complete checklist of needed documents
        🔗 **Official Resources**: FBR website links and helpline numbers
        🚨 **Common Mistakes**: What to avoid and red flags
        💡 **Expert Tips**: Professional advice and best practices
        📞 **Next Steps**: Clear action items prioritized by urgency
        
        IMPORTANT: Focus specifically on Pakistani tax law, FBR regulations, and current tax year requirements. 
        Always recommend verifying current information with FBR official sources.
        Make it practical, actionable, and professionally formatted.
        """
        
        completion = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1500,
        )
        
        return completion.choices[0].message.content
            
    except Exception as e:
        return f"Error processing your tax question: {str(e)}"

def chatbot_response(user_query):
    """Enhanced general chatbot with better tax redirection"""
    if not groq_client:
        return "Hello! I'm a general chatbot but I need an API key to function properly."
    
    try:
        prompt = f"""
        You are a friendly, general-purpose chatbot assistant. Keep responses conversational and helpful.
        
        User query: "{user_query}"
        
        IMPORTANT: If this query is related to Pakistani taxes, FBR, tax forms, or tax calculations, respond with: 
        "I notice you're asking about Pakistani taxes! For expert guidance on tax forms, filing procedures, calculations, and FBR regulations, please use the Tax Assistant feature above. I specialize in general conversations, while the Tax Assistant provides comprehensive Pakistani tax expertise."
        
        For all other topics, provide a helpful, warm, conversational response. Keep it concise and natural.
        """
        
        completion = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=200,
        )
        
        return completion.choices[0].message.content
    except Exception as e:
        return f"Sorry, I encountered an error: {str(e)}"

# ─────────────────────────────────────────
# 📚 Enhanced Utility Functions
# ─────────────────────────────────────────
def add_to_history(country, query, pdf_url=None, fillable_fields=0):
    """Add search to history"""
    if "search_history" not in st.session_state:
        st.session_state.search_history = []
    
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.search_history.append({
        "timestamp": timestamp,
        "country": country,
        "query": query,
        "pdf_url": pdf_url,
        "fillable_fields": fillable_fields
    })
    
    # Keep only last 15 searches
    if len(st.session_state.search_history) > 15:
        st.session_state.search_history = st.session_state.search_history[-15:]

def initialize_session_state():
    """Initialize session state variables for each tab independently"""
    defaults = {
        # Global states
        "search_history": [],
        "active_tab": "🤖 Smart Assistant",
        "assistant_mode": "tax",  # 'tax' or 'chat'
        
        # Assistant tab states  
        "assistant_responses": [],
        "last_tax_type": None,
        
        # Find Forms tab states
        "find_forms_results": [],
        "find_forms_query": "",
        "find_forms_executed": False,
        
        # Upload Forms tab states
        "upload_pdf_bytes": None,
        "upload_form_fields": [],
        "upload_file_analysis": None,
        "upload_file_type": None,
    }
    
    for key, default_value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = default_value

def get_tax_specific_actions(tax_type):
    """Get relevant quick actions based on identified tax type"""
    tax_actions = {
        'income_tax': [
            ("📄 Income Tax Guide", "Complete guide for filing income tax returns in Pakistan"),
            ("🧮 Tax Calculator", "Calculate your income tax liability and advance tax"),
            ("📅 Filing Deadlines", "Income tax return filing deadlines and due dates")
        ],
        'sales_tax': [
            ("📊 Sales Tax Guide", "Sales tax registration and return filing procedures"),
            ("🧾 GST Compliance", "Goods and Services Tax compliance requirements"),
            ("📋 Sales Tax Returns", "How to file monthly sales tax returns")
        ],
        'withholding_tax': [
            ("✂️ Withholding Tax Guide", "Withholding tax deduction rates and procedures"),
            ("📑 WHT Statements", "Filing withholding tax statements and certificates"),
            ("💳 Advance Tax", "Advance tax payment procedures and calculations")
        ],
        'property_tax': [
            ("🏠 Property Tax Guide", "Property tax assessment and payment procedures"),
            ("📋 Property Registration", "Property registration tax and documentation"),
            ("💰 Capital Gains Tax", "Capital gains tax on property transactions")
        ],
        'wealth_tax': [
            ("💎 Wealth Statement", "Filing wealth statements and asset declarations"),
            ("🏦 Asset Declaration", "Declaring foreign and domestic assets"),
            ("📊 Wealth Tax Compliance", "Wealth tax compliance requirements")
        ],
        'business_tax': [
            ("🏢 Business Tax Guide", "Corporate tax and business registration procedures"),
            ("📈 Company Returns", "Filing company tax returns and compliance"),
            ("⚖️ Tax Audit", "Tax audit procedures and requirements for businesses")
        ],
        'general_tax': [
            ("📋 Tax Overview", "Overview of Pakistani tax system and types"),
            ("🆔 NTN Registration", "National Tax Number registration procedures"),
            ("📞 FBR Services", "FBR contact information and services")
        ]
    }
    
    return tax_actions.get(tax_type, tax_actions['general_tax'])

def get_quick_actions_response(action_type, tax_guide_type):
    """Enhanced quick action responses based on tax type"""
    if not groq_client:
        return f"Quick action: {action_type}. Please configure Groq API for detailed guidance."
    
    try:
        prompt = f"""
        You are LifePilot Tax Agent providing expert guidance for Pakistani taxation.
        
        Action Type: "{action_type}"
        Tax Category: "{tax_guide_type}"
        
        Provide comprehensive, actionable guidance for this specific Pakistani tax matter.
        Include relevant forms, procedures, deadlines, rates, and FBR requirements.
        Structure with clear headers, emojis, and step-by-step instructions where applicable.
        Focus on current tax year requirements and always mention to verify with FBR.
        
        Make it practical and implementable for Pakistani taxpayers.
        """
        
        completion = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1000,
        )
        
        return completion.choices[0].message.content
    except Exception as e:
        return f"Error getting guidance for {action_type}: {str(e)}"

# ─────────────────────────────────────────
# 🏠 Main Application
# ─────────────────────────────────────────
def main():
    st.set_page_config(
        page_title="LifePilot - Pakistan Tax Forms", 
        page_icon="📋",
        layout="wide"
    )
    st.markdown("""
    <style>
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    }
    .main .block-container {
        background: rgba(255, 255, 255, 0.9);
        border-radius: 10px;
        padding: 2rem;
        margin-top: 1rem;
    }
    </style>
    """, unsafe_allow_html=True)
        
    initialize_session_state()

    # Sidebar
    with st.sidebar:
        st.title("📋 LifePilot")
        st.caption("Pakistan Tax Assistant")

        st.subheader("📊 Recent Searches")
        if st.session_state.search_history:
            for item in reversed(st.session_state.search_history[-5:]):
                with st.expander(f"🔍 {item['query'][:25]}..."):
                    st.write(f"📅 {item['timestamp']}")
                    st.write(f"📝 Fields: {item.get('fillable_fields', 0)}")
        else:
            st.info("Search history will appear here")

        st.divider()
        
        # API Status
        st.subheader("🔧 System Status")
        st.write("🔍 Search:", "✅ Active" if SERPER_API_KEY else "⚠️ Limited")
        st.write("🤖 AI Agent:", "✅ Active" if GROQ_API_KEY else "⚠️ Limited")
        
        st.divider()
        st.markdown("**📋 LifePilot** | Pakistan Tax Expert")
        st.markdown("*Powered by AI • FBR Compliant*")

    # Main content
    st.title("🇵🇰 Pakistan Tax Assistant")
    st.markdown("Your intelligent assistant for Pakistani tax forms, filing procedures, and FBR regulations.")
    
    # Create tabs using buttons in columns
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("🤖 Smart Assistant", use_container_width=True, 
                    type="primary" if st.session_state.active_tab == "🤖 Smart Assistant" else "secondary"):
            st.session_state.active_tab = "🤖 Smart Assistant"
    
    with col2:
        if st.button("🔍 Find Forms", use_container_width=True,
                    type="primary" if st.session_state.active_tab == "🔍 Find Forms" else "secondary"):
            st.session_state.active_tab = "🔍 Find Forms"
    
    with col3:
        if st.button("📤 Upload Forms", use_container_width=True,
                    type="primary" if st.session_state.active_tab == "📤 Upload Forms" else "secondary"):
            st.session_state.active_tab = "📤 Upload Forms"
    
    st.markdown("---")
    
    # Display content based on active tab
    if st.session_state.active_tab == "🤖 Smart Assistant":
        st.header("🧠 Advanced Tax Assistant")
        st.markdown("Get expert guidance on Pakistani taxes, forms, filing procedures, and FBR regulations.")
        
        # Mode selection buttons
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🎯 Tax Assistant", use_container_width=True, 
                        type="primary" if st.session_state.assistant_mode == "tax" else "secondary"):
                st.session_state.assistant_mode = "tax"
        
        with col2:
            if st.button("💬 General Chatbot", use_container_width=True,
                        type="primary" if st.session_state.assistant_mode == "chat" else "secondary"):
                st.session_state.assistant_mode = "chat"
        
        st.markdown("---")
        
        # Input based on selected mode
        if st.session_state.assistant_mode == "tax":
            user_query = st.text_input(
                "🎯 Ask your Pakistani tax question:", 
                placeholder="Which form for income tax? How to calculate advance tax? Sales tax registration procedure?",
                key="tax_assistant_query"
            )
        else:
            user_query = st.text_input(
                "💬 Chat with me about anything:", 
                placeholder="How are you? Tell me a joke, or ask about general topics...",
                key="general_chat_query"
            )
        
        if user_query:
            if st.session_state.assistant_mode == "tax":
                # Tax Assistant Mode
                tax_type = identify_tax_type(user_query)
                st.session_state.last_tax_type = tax_type
                
                with st.spinner("🧠 Analyzing your tax question..."):
                    response = tax_agent_response(
                        user_query, 
                        tax_form_type=st.session_state.get('upload_file_type'),
                        form_fields=st.session_state.upload_form_fields
                    )
                
                st.success("**🎯 Tax Expert Analysis:**")
                st.markdown(response)
                
                # Add to assistant responses
                st.session_state.assistant_responses.append({
                    "query": user_query,
                    "response": response,
                    "tax_type": tax_type,
                    "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
                })
                
                # Show tax-specific quick actions
                st.markdown("### 🚀 Related Tax Guidance:")
                tax_actions = get_tax_specific_actions(tax_type)
                
                col1, col2, col3 = st.columns(3)
                for idx, (action_name, action_desc) in enumerate(tax_actions):
                    with [col1, col2, col3][idx % 3]:
                        if st.button(action_name, key=f"tax_action_{idx}_{tax_type}", help=action_desc):
                            with st.spinner(f"Getting {action_name.lower()} guidance..."):
                                guide_response = get_quick_actions_response(action_name, tax_type)
                            st.markdown(f"### {action_name}")
                            st.markdown(guide_response)
            else:
                # General Chatbot Mode
                with st.spinner("💭 Thinking..."):
                    response = chatbot_response(user_query)
                st.info("**💬 General Response:**")
                st.markdown(response)

    elif st.session_state.active_tab == "🔍 Find Forms":
        st.header("🔍 Find Tax Forms")
        st.markdown("Search and download official Pakistani tax forms from FBR and government sources.")
        
        col1, col2 = st.columns([2, 1])
        with col1:
            form_type = st.selectbox(
                "Select form type:",
                ["Income Tax Return", "Sales Tax Return", "Withholding Tax Statement", 
                 "Property Tax", "Advance Tax", "Wealth Statement", "Business Registration",
                 "NTN Application", "Tax Exemption Certificate", "Custom Forms"],
                index=0
            )
        
        with col2:
            details = st.text_input("Additional details:", 
                                  placeholder="individual, salaried, business")
        
        search_query = f"{form_type} {details}".strip()
        
        # Enhanced search options
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("🔍 Search Forms", use_container_width=True, type="primary"):
                if search_query != st.session_state.find_forms_query or not st.session_state.find_forms_executed:
                    st.session_state.find_forms_query = search_query
                    st.session_state.find_forms_executed = True
                    
                    with st.spinner("🔎 Searching for Pakistani tax forms..."):
                        results = serper_search(search_query, "pk")
                        st.session_state.find_forms_results = results
        
        with col2:
            if st.button("🗑️ Clear Results", use_container_width=True):
                st.session_state.find_forms_results = []
                st.session_state.find_forms_executed = False
                st.session_state.find_forms_query = ""
                if 'find_forms_pdf' in st.session_state:
                    del st.session_state.find_forms_pdf
                if 'find_forms_fields' in st.session_state:
                    del st.session_state.find_forms_fields
                st.success("Results cleared!")
                        
        # Display stored results for Find Forms tab only
        if st.session_state.find_forms_results and st.session_state.find_forms_executed:
            results = st.session_state.find_forms_results
            st.success(f"✅ Found {len(results)} relevant tax forms")
            
            for idx, result in enumerate(results[:6]):  # Show up to 6 results
                with st.container():
                    st.subheader(f"{idx+1}. {result.get('title', 'Tax Form')}")
                    st.write(result.get('snippet', ''))
                    
                    col1, col2, col3 = st.columns([1, 1, 1])
                    with col1:
                        # Download button for all forms (fillable or not)
                        if st.button(f"📥 Download", key=f"download_form_find_{idx}"):
                            with st.spinner("📄 Downloading form..."):
                                pdf_bytes = fetch_pdf(result.get('link', ''))
                                if pdf_bytes:
                                    # Check if fillable
                                    is_fillable, field_count = is_pdf_fillable(pdf_bytes)
                                    
                                    st.session_state.find_forms_pdf = pdf_bytes
                                    st.session_state.find_forms_fields = extract_form_fields(pdf_bytes) if is_fillable else []
                                    add_to_history("Pakistan", search_query, result.get('link'), field_count)
                                    
                                    if is_fillable:
                                        st.success(f"✅ Fillable form downloaded! ({field_count} fields)")
                                    else:
                                        st.success("✅ Form downloaded successfully!")
                                        st.info("📄 This form is not fillable but can be printed and filled manually")
                                else:
                                    st.error("❌ Could not download form from this link")
                    
                    with col2:
                        if st.button(f"👁️ Preview", key=f"preview_form_find_{idx}"):
                            with st.spinner("Loading preview..."):
                                pdf_bytes = fetch_pdf(result.get('link', ''))
                                if pdf_bytes:
                                    st.session_state.find_forms_pdf = pdf_bytes
                                    is_fillable, field_count = is_pdf_fillable(pdf_bytes)
                                    st.session_state.find_forms_fields = extract_form_fields(pdf_bytes) if is_fillable else []
                                    st.success("✅ Form loaded for preview")
                                else:
                                    st.error("❌ Could not load preview")
                    
                    with col3:
                        st.markdown(f"[🔗 Source]({result.get('link', '#')})")
                    
                    st.divider()
        
        elif st.session_state.find_forms_executed and not st.session_state.find_forms_results:
            st.warning("⚠️ No forms found. Try different keywords or check your internet connection.")
        
        # Display downloaded form in Find Forms tab
        if 'find_forms_pdf' in st.session_state and st.session_state.find_forms_pdf:
            st.markdown("---")
            st.markdown("## 📋 Downloaded Form")
            
            field_count = len(st.session_state.get('find_forms_fields', []))
            if field_count > 0:
                st.info(f"📄 **Fillable PDF Form** • {field_count} interactive fields detected")
            else:
                st.info("📄 **PDF Form** • Static form (print and fill manually)")
            
            tab1, tab2 = st.tabs(["📖 PDF Preview", "📝 Form Analysis"])
            
            with tab1:
                display_pdf(st.session_state.find_forms_pdf)
            
            with tab2:
                if st.session_state.get('find_forms_fields', []):
                    st.markdown(f"### 📝 Interactive Form Fields ({len(st.session_state.find_forms_fields)})")
                    
                    if len(st.session_state.find_forms_fields) > 0:
                        df_fields = pd.DataFrame(st.session_state.find_forms_fields)
                        df_fields = df_fields[['name', 'type', 'page']].head(25)
                        st.dataframe(df_fields, use_container_width=True)
                        
                        if len(st.session_state.find_forms_fields) > 25:
                            st.info(f"Showing first 25 fields out of {len(st.session_state.find_forms_fields)} total")
                
                # Dynamic LLM Analysis
                st.markdown("### 🤖 AI Form Analysis")
                with st.spinner("🧠 Analyzing form with AI..."):
                    analysis = analyze_form_with_llm(
                        file_type='pdf',
                        form_fields=st.session_state.get('find_forms_fields', []),
                        search_query=st.session_state.get('find_forms_query', '')
                    )
                st.markdown(analysis)
                        
                # Form filling assistance
                st.markdown("### ❓ Need help filling this form?")
                if st.button("🤖 Get Form Filling Guide", key="form_filling_help"):
                    with st.spinner("Generating form filling guidance..."):
                        form_help = tax_agent_response(f"Help me fill this tax form with {field_count} fields")
                    st.markdown("### 📋 Form Filling Guide")
                    st.markdown(form_help)

    elif st.session_state.active_tab == "📤 Upload Forms":
        st.header("📤 Upload Tax Forms")
        st.markdown("Upload your tax forms for analysis, assistance, and form-specific guidance")
        
        uploaded_file = st.file_uploader(
            "Choose your tax form:", 
            type=["pdf", "xlsx", "xls", "docx", "doc"],
            help="Upload PDF, Excel, or Word tax forms for analysis"
        )
        
        if uploaded_file:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            file_bytes = BytesIO(uploaded_file.read())
            
            with st.spinner("📊 Analyzing your tax form..."):
                if file_extension == 'pdf':
                    is_fillable, field_count = is_pdf_fillable(file_bytes)
                    
                    st.session_state.upload_pdf_bytes = file_bytes
                    st.session_state.upload_form_fields = extract_form_fields(file_bytes) if is_fillable else []
                    st.session_state.upload_file_type = 'pdf'
                    
                    if is_fillable:
                        st.success(f"✅ Fillable PDF uploaded! Found {field_count} interactive fields.")
                    else:
                        st.success("✅ PDF uploaded successfully! This is a static form.")
                        st.info("📄 Static form detected - can be printed and filled manually")
                        
                elif file_extension in ['xlsx', 'xls']:
                    analysis = analyze_excel_file(file_bytes)
                    if analysis:
                        st.session_state.upload_file_analysis = analysis
                        st.session_state.upload_file_type = 'excel'
                        if analysis['is_form_like']:
                            st.success("✅ Excel tax form uploaded successfully!")
                        else:
                            st.success("✅ Excel file uploaded successfully!")
                            st.info("📊 This may not be a standard tax form, but I can still help analyze it")
                    else:
                        st.error("❌ Could not analyze Excel file")
                        
                elif file_extension in ['docx', 'doc']:
                    analysis = analyze_word_file(file_bytes)
                    if analysis:
                        st.session_state.upload_file_analysis = analysis
                        st.session_state.upload_file_type = 'word'
                        if analysis['is_form_like']:
                            st.success("✅ Word tax form uploaded successfully!")
                        else:
                            st.success("✅ Word document uploaded successfully!")
                            st.info("📝 This may not be a standard tax form, but I can still help analyze it")
                    else:
                        st.error("❌ Could not analyze Word document")

        # Display current form section for Upload tab only
        if st.session_state.upload_pdf_bytes or st.session_state.upload_file_analysis:
            st.markdown("---")
            st.markdown("## 📋 Uploaded Tax Form Analysis")
            
            # Form information header
            file_type = st.session_state.upload_file_type
            if file_type == 'pdf':
                field_count = len(st.session_state.upload_form_fields)
                if field_count > 0:
                    st.info(f"📄 **Fillable PDF Form** • {field_count} interactive fields detected")
                else:
                    st.info("📄 **Static PDF Form** • No interactive fields (print and fill manually)")
            elif file_type == 'excel':
                field_count = st.session_state.upload_file_analysis.get('total_potential_fields', 0)
                sheet_count = len(st.session_state.upload_file_analysis.get('sheets', []))
                st.info(f"📊 **Excel Form** • {sheet_count} sheets • {field_count} potential form fields")
            elif file_type == 'word':
                field_count = st.session_state.upload_file_analysis.get('form_fields_found', 0)
                st.info(f"📝 **Word Document** • {field_count} form elements detected")
            
            # Action buttons
            col1, col2 = st.columns([1, 4])
            with col1:
                if st.button("🗑️ Clear Form", type="secondary"):
                    # Clear upload tab form-related session state only
                    keys_to_clear = ['upload_pdf_bytes', 'upload_form_fields', 'upload_file_analysis', 'upload_file_type']
                    for key in keys_to_clear:
                        if key == 'upload_form_fields':
                            st.session_state[key] = []
                        else:
                            st.session_state[key] = None
                    st.success("✅ Form cleared!")
                    st.experimental_rerun()
            
            # Display form content based on type
            if file_type == 'pdf' and st.session_state.upload_pdf_bytes:
                tab1, tab2 = st.tabs(["📖 PDF Preview", "📝 Form Analysis"])
                
                with tab1:
                    display_pdf(st.session_state.upload_pdf_bytes)
                
                with tab2:
                    field_count = len(st.session_state.upload_form_fields)
                    
                    if field_count > 0:
                        st.markdown(f"### 📝 Interactive Form Fields ({field_count})")
                        
                        # Create a DataFrame for better display
                        df_fields = pd.DataFrame(st.session_state.upload_form_fields)
                        df_fields = df_fields[['name', 'type', 'page']].head(30)
                        st.dataframe(df_fields, use_container_width=True)
                        
                        if field_count > 30:
                            st.info(f"Showing first 30 fields out of {field_count} total")
                    
                    # Dynamic LLM Analysis
                    st.markdown("### 🤖 AI Form Analysis")
                    with st.spinner("🧠 Analyzing uploaded form with AI..."):
                        analysis = analyze_form_with_llm(
                            file_type=file_type,
                            form_fields=st.session_state.upload_form_fields,
                            file_analysis=st.session_state.upload_file_analysis
                        )
                    st.markdown(analysis)
            
            elif file_type in ['excel', 'word'] and st.session_state.upload_file_analysis:
                analysis = st.session_state.upload_file_analysis
                
                st.markdown("### 📊 Form Structure Analysis")
                
                if file_type == 'excel':
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("Total Sheets", len(analysis['sheets']))
                        st.metric("Potential Form Fields", analysis['total_potential_fields'])
                    
                    with col2:
                        st.markdown("**Sheet Details:**")
                        for sheet in analysis['sheets']:
                            st.write(f"• **{sheet['name']}**: {sheet['rows']}×{sheet['columns']} ({sheet['potential_fields']} fields)")
                
                else:  # word
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Paragraphs", analysis['paragraphs'])
                    with col2:
                        st.metric("Tables", analysis['tables'])
                    with col3:
                        st.metric("Form Elements", analysis['form_fields_found'])
                
                # Dynamic LLM Analysis
                st.markdown("### 🤖 AI Form Analysis")
                with st.spinner("🧠 Analyzing uploaded form with AI..."):
                    llm_analysis = analyze_form_with_llm(
                        file_type=file_type,
                        file_analysis=st.session_state.upload_file_analysis
                    )
                st.markdown(llm_analysis)
            
            # Enhanced form assistance section
            st.markdown("### ❓ Form-Specific Assistance")
            
            # Quick help buttons
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("📋 How to Fill", key="how_to_fill", use_container_width=True):
                    with st.spinner("Getting form filling guidance..."):
                        response = tax_agent_response(f"How do I fill this {file_type} tax form step by step?")
                    st.markdown("### 📋 Form Filling Guide")
                    st.markdown(response)
            
            with col2:
                if st.button("📄 Required Documents", key="required_docs", use_container_width=True):
                    with st.spinner("Getting document requirements..."):
                        response = tax_agent_response("What documents do I need to complete this tax form?")
                    st.markdown("### 📄 Required Documents")
                    st.markdown(response)
            
            with col3:
                if st.button("⚠️ Common Mistakes", key="common_mistakes", use_container_width=True):
                    with st.spinner("Getting common mistakes info..."):
                        response = tax_agent_response("What are common mistakes to avoid when filling this tax form?")
                    st.markdown("### ⚠️ Common Mistakes to Avoid")
                    st.markdown(response)
            
            # Custom question section
            col1, col2 = st.columns([3, 1])
            with col1:
                help_query = st.text_input(
                    "Ask about this specific form:", 
                    placeholder="How to calculate section B? What if I have no income? How to handle deductions?",
                    key="form_help_query"
                )
            
            with col2:
                ask_help = st.button("🤖 Ask Expert", type="primary", use_container_width=True)
            
            if help_query and ask_help:
                with st.spinner("🧠 Analyzing your form-specific question..."):
                    response = tax_agent_response(
                        help_query, 
                        st.session_state.upload_file_type, 
                        st.session_state.upload_form_fields
                    )
                    
                st.markdown("### 🎯 Expert Form Guidance")
                st.success("**Personalized Form Assistance:**")
                st.markdown(response)

    # Enhanced Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666; padding: 20px;'>
        <p><strong>📋 LifePilot Tax Assistant</strong> | Powered by Advanced AI | Specialized for Pakistani Taxation</p>
        <p style='font-size: 0.9em;'>✅ FBR Compliant Information • 🔒 Secure & Private • 🇵🇰 Pakistan-Focused</p>
        <p style='font-size: 0.8em; color: #888;'>This tool provides general tax guidance. For complex matters, consult qualified tax professionals or contact FBR directly.</p>
        <p style='font-size: 0.8em;'><strong>FBR Helpline:</strong> 111-772-772 | <strong>Website:</strong> fbr.gov.pk</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()

